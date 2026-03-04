"""
Script to generate the AIT CMMS Program Overview Word Document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_bg(cell, color_hex):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    heading.paragraph_format.space_before = Pt(14)
    heading.paragraph_format.space_after = Pt(6)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(3)
    return para

def add_kpi_table(doc, kpis):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    headers = ['KPI Name', 'What It Measures', 'Business Value']
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr[i], '1F3864')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(kpis):
        row = table.add_row().cells
        for j, val in enumerate(row_data):
            row[j].text = val
            row[j].paragraphs[0].runs[0].font.size = Pt(10)
            if i % 2 == 0:
                set_cell_bg(row[j], 'DCE6F1')
            else:
                set_cell_bg(row[j], 'FFFFFF')
    return table

def add_cost_table(doc, rows, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr[i], '1F3864')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(rows):
        row = table.add_row().cells
        for j, val in enumerate(row_data):
            row[j].text = val
            row[j].paragraphs[0].runs[0].font.size = Pt(10)
            if i % 2 == 0:
                set_cell_bg(row[j], 'E2EFDA')
            else:
                set_cell_bg(row[j], 'FFFFFF')
    return table


def build_document():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Default body font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════
    # COVER / TITLE BLOCK
    # ══════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('AIT CMMS 2.3.1')
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(31, 56, 100)
    title_run.font.name = 'Calibri'

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run('Computerized Maintenance Management System')
    sub_run.bold = True
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()

    tag_para = doc.add_paragraph()
    tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag_para.add_run(
        'Program Overview & Financial Value Analysis\n'
        'Prepared for AIT Organization Leadership'
    )
    tag_run.italic = True
    tag_run.font.size = Pt(12)
    tag_run.font.color.rgb = RGBColor(80, 80, 80)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.date.today().strftime('Date: %B %d, %Y'))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ══════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ══════════════════════════════════════════
    add_heading(doc, '1. Executive Summary', level=1, color=(31, 56, 100))

    doc.add_paragraph(
        'AIT CMMS 2.3.1 is a purpose-built, enterprise-grade Computerized Maintenance '
        'Management System developed exclusively for AIT\'s operational environment. The '
        'system replaces fragmented spreadsheets, paper-based work orders, and manual '
        'tracking with a single, centralized digital platform that automates preventive '
        'maintenance scheduling, corrective work order management, inventory control, and '
        'key performance indicator (KPI) reporting.'
    )

    doc.add_paragraph(
        'By digitizing every aspect of the maintenance lifecycle, AIT CMMS 2.3.1 directly '
        'reduces equipment downtime, lowers emergency repair costs, prevents excess inventory '
        'spend, and generates the auditable data needed to demonstrate maintenance excellence '
        'to leadership, auditors, and customers. The financial impact compounds over time: '
        'savings realized in Year 1 from reduced breakdowns and optimized labor continue to '
        'grow as the system\'s historical database grows richer and scheduling becomes '
        'increasingly precise.'
    )

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # HOW THE PROGRAM WORKS
    # ══════════════════════════════════════════
    add_heading(doc, '2. How the Program Works', level=1, color=(31, 56, 100))

    doc.add_paragraph(
        'AIT CMMS 2.3.1 is a multi-user desktop application written in Python 3 with a '
        'PostgreSQL relational database backend. Technicians and supervisors access the '
        'system simultaneously over the local network; every action is recorded with a '
        'timestamp and the responsible user\'s name, providing a complete audit trail.'
    )

    # 2.1 Architecture
    add_heading(doc, '2.1  System Architecture', level=2, color=(70, 130, 180))
    arch_items = [
        ('Frontend (User Interface)', 'Python / Tkinter – lightweight, cross-platform '
         'graphical interface that runs on existing Windows workstations with no browser or '
         'web server required.'),
        ('Backend Database', 'PostgreSQL 18.1 installed on a local server – enterprise-class '
         'reliability with zero recurring licensing cost.'),
        ('Connection Layer', 'Threaded connection pooling with TCP keepalive ensures all '
         'concurrent users share resources efficiently without dropped connections.'),
        ('Security', 'Username/password authentication, bcrypt password hashing, role-based '
         'access, and an immutable audit log of every change.'),
        ('Reporting Engine', 'ReportLab for PDF generation; pandas for data aggregation; '
         'python-docx for Word exports – all producing professional, print-ready documents.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Layer'
    hdr[1].text = 'Technology & Role'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(cell, '1F3864')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (layer, desc) in enumerate(arch_items):
        row = table.add_row().cells
        row[0].text = layer
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = desc
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
        bg = 'DCE6F1' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
    doc.add_paragraph()

    # 2.2 Core Modules
    add_heading(doc, '2.2  Core Functional Modules', level=2, color=(70, 130, 180))

    # Preventive Maintenance
    add_heading(doc, '2.2.1  Preventive Maintenance (PM) Scheduling Engine', level=3)
    doc.add_paragraph(
        'The PM Scheduling Engine is the operational heart of AIT CMMS. It automatically '
        'determines which equipment is due for preventive maintenance across four scheduling '
        'frequencies — Weekly, Monthly, 6-Month, and Annual — and produces optimized daily '
        'work lists for technicians.'
    )
    for item in [
        'Reads every active asset from the equipment register and computes its next due date based on completion history.',
        'Applies conflict-resolution logic to distribute PM load evenly and avoid technician bottlenecks.',
        'Assigns tasks to available technicians by skill set and current workload.',
        'Automatically flags overdue PMs so nothing falls through the cracks.',
        'Records labor hours, technician name, parts used, and completion notes for every finished PM.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # Corrective Maintenance
    add_heading(doc, '2.2.2  Corrective Maintenance (CM) Work Order Management', level=3)
    doc.add_paragraph(
        'When equipment fails or a defect is discovered, technicians create a CM work order '
        'inside AIT CMMS. The system manages the full lifecycle:'
    )
    for item in [
        'Priority classification (Critical, High, Medium, Low) with visual status tracking.',
        'Automatic parts-request creation — technicians specify what parts they need; the system checks MRO stock and flags shortages.',
        'Labor and cost capture: every hour spent and every part consumed is linked to the work order.',
        'Historical CM records feed directly into MTBF (Mean Time Between Failures) KPI calculations, revealing which assets are chronic problem generators.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # Equipment Management
    add_heading(doc, '2.2.3  Equipment Register & Asset Management', level=3)
    doc.add_paragraph(
        'Every maintainable asset at AIT is stored in a structured equipment register with '
        'full metadata: BFM number, description, location, photo, PM frequencies, and current '
        'status (Active / Run-to-Failure / Missing / Deactivated). The register is the single '
        'source of truth that drives all scheduling, reporting, and KPI calculations.'
    )
    doc.add_paragraph()

    # MRO Inventory
    add_heading(doc, '2.2.4  MRO Inventory & Parts Management', level=3)
    doc.add_paragraph(
        'The MRO (Maintenance, Repair, and Operations) Stock Module tracks every spare part '
        'and consumable used by the maintenance department:'
    )
    for item in [
        'Part number, description, unit cost, quantity on hand, minimum stock threshold, supplier, and lead time.',
        'Physical bin/rack/row location for rapid retrieval.',
        'Automatic low-stock alerts prevent stockouts that would delay repairs.',
        'Cost calculations (Qty × Unit Price) flow automatically into work order cost reports.',
        'Photo storage helps technicians identify the correct part quickly.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # KPI
    add_heading(doc, '2.2.5  KPI Management & Analytics Dashboard', level=3)
    doc.add_paragraph(
        'AIT CMMS includes a dedicated 2025 KPI framework with 17+ predefined indicators. '
        'The system auto-collects data from completed work orders and PM records, calculates '
        'results quarterly, and presents trend charts and target-vs-actual comparisons in a '
        'dashboard view. KPI results can be exported to PDF or Excel for board-level reporting.'
    )
    doc.add_paragraph()

    kpis = [
        ('PM Adherence Rate',        'Percentage of PMs completed on schedule',       'Demonstrates maintenance discipline; directly linked to equipment reliability'),
        ('MTBF',                     'Mean Time Between Failures per asset',           'Identifies unreliable assets consuming disproportionate repair budget'),
        ('TTR Adherence',            'Actual vs. target Time-to-Repair',              'Measures technician efficiency and parts availability'),
        ('Technical Availability',   'Uptime percentage of critical equipment',        'Quantifies production capacity protected by maintenance'),
        ('Work Order Backlog',        'Open WOs vs. total WO volume',                  'Reveals capacity gaps before they become crises'),
        ('Mean Response Time (MRT)', 'Time from fault report to technician on-site',  'Measures urgency-response alignment; impacts safety outcomes'),
        ('Injury Frequency Rate',    'Safety incidents per 200,000 hours worked',      'Demonstrates safety culture; impacts insurance and compliance costs'),
        ('Cost per Maintenance Event','Average cost of PM and CM activities',         'Benchmarks cost efficiency over time and vs. industry'),
    ]
    add_kpi_table(doc, kpis)
    doc.add_paragraph()

    # Reporting
    add_heading(doc, '2.2.6  Reporting & Document Management', level=3)
    doc.add_paragraph(
        'AIT CMMS generates a range of professional, print-ready reports without any manual '
        'data assembly:'
    )
    for item in [
        'Weekly PM completion reports for supervisors.',
        'Monthly maintenance summary reports with cost totals and labor hours.',
        'Quarterly KPI reports with trend analysis for leadership review.',
        'Equipment maintenance history exports per asset.',
        'MRO inventory valuation and stock-level reports.',
        'Audit logs for compliance and internal review.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # Multi-User
    add_heading(doc, '2.2.7  Multi-User Access & Security', level=3)
    doc.add_paragraph(
        'Any number of technicians and supervisors can log in simultaneously from different '
        'workstations on the local network. Optimistic concurrency control prevents two users '
        'from overwriting each other\'s edits. Role-based permissions ensure that only '
        'authorized personnel can modify sensitive records or administrative settings. Every '
        'create, update, and delete action is recorded in an immutable audit log with the '
        'user\'s name, the field changed, the old value, the new value, and the timestamp.'
    )
    doc.add_paragraph()

    # Backup
    add_heading(doc, '2.2.8  Automated Backup & Data Protection', level=3)
    doc.add_paragraph(
        'The integrated Backup Manager runs automated daily pg_dump snapshots of the entire '
        'database, retains 30 days of rolling backups, and provides a one-click restore '
        'procedure. This ensures that even in a catastrophic hardware failure, AIT loses at '
        'most one day of data.'
    )
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════
    # FINANCIAL ADVANTAGES
    # ══════════════════════════════════════════
    add_heading(doc, '3. Financial Advantages for AIT Organization', level=1, color=(31, 56, 100))

    doc.add_paragraph(
        'The financial case for AIT CMMS 2.3.1 rests on five pillars: cost avoidance through '
        'equipment reliability, labor efficiency, inventory optimization, elimination of '
        'external software costs, and regulatory/compliance risk reduction. Each is described '
        'below with illustrative figures based on industry benchmarks for maintenance '
        'operations of AIT\'s scale.'
    )

    # 3.1 Reduced Unplanned Downtime
    add_heading(doc, '3.1  Reduced Unplanned Downtime & Emergency Repair Costs', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'Unplanned equipment failures are the most expensive maintenance events. Emergency '
        'labor (often overtime or contracted specialists), expedited parts shipping, and lost '
        'production capacity combine to make reactive maintenance 3–5× more expensive per '
        'event than a planned preventive maintenance task.'
    )
    doc.add_paragraph(
        'AIT CMMS eliminates the primary cause of unplanned breakdowns — missed or forgotten '
        'PMs — by automating the scheduling and tracking of every preventive task. When PM '
        'adherence improves from a typical manual baseline of 60–70% to 90%+, the '
        'frequency of unexpected failures falls proportionally.'
    )

    cost_rows_downtime = [
        ('Scenario',                        'Manual / Spreadsheet System',   'AIT CMMS (Automated)'),
        ('PM Adherence Rate',               '60–70%',                        '90–95%'),
        ('Unplanned Failures per Quarter',  '15–20 events (estimated)',       '4–8 events (estimated)'),
        ('Avg. Cost per Emergency Event',   '$800–$2,500',                   '$800–$2,500'),
        ('Quarterly Emergency Repair Cost', '$12,000–$50,000',               '$3,200–$20,000'),
        ('Estimated Annual Savings',        '—',                             '$35,000–$120,000+'),
    ]
    headers_dt = ['Metric', 'Without CMMS', 'With AIT CMMS']
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(headers_dt):
        hdr2[i].text = h
        hdr2[i].paragraphs[0].runs[0].bold = True
        hdr2[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr2[i], '1F3864')
        hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr2[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (metric, without, with_cmms) in enumerate(cost_rows_downtime):
        row = table2.add_row().cells
        row[0].text = metric
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = without
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
        row[2].text = with_cmms
        row[2].paragraphs[0].runs[0].font.size = Pt(10)
        bg = 'E8F5E9' if i % 2 == 0 else 'FFFFFF'
        for cell in row:
            set_cell_bg(cell, bg)
    doc.add_paragraph()

    # 3.2 Labor Efficiency
    add_heading(doc, '3.2  Improved Labor Efficiency & Technician Productivity', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'Without a CMMS, technicians and supervisors spend significant time each week on '
        'non-wrench activities: searching paper logs for asset history, manually building '
        'PM schedules in spreadsheets, looking up part locations, and compiling reports. '
        'Industry studies consistently show that a well-implemented CMMS recaptures '
        '15–25% of a technician\'s productive time.'
    )

    labor_rows = [
        ('Supervisors searching/building schedules',     '5–8 hrs/week',   '< 1 hr/week (auto-generated)'),
        ('Technicians locating parts in storeroom',      '30–60 min/event','5–10 min (bin location in system)'),
        ('Report compilation (monthly/quarterly)',       '4–6 hrs/report', '< 30 min (auto-generated)'),
        ('Root-cause research (no history available)',   '2–4 hrs/event',  '< 15 min (full history in system)'),
        ('Total non-productive time saved per week',     '—',              '8–16 hrs (est.)'),
    ]
    add_cost_table(doc,
        labor_rows,
        ['Activity', 'Time Without CMMS', 'Time With AIT CMMS']
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'At an average fully-loaded technician cost of $35–$55/hour, recovering 10 hours '
        'per week across the maintenance team represents $18,000–$28,600 per year in '
        'productive labor that can be redirected to value-adding maintenance work rather '
        'than administrative overhead.'
    )
    doc.add_paragraph()

    # 3.3 Inventory
    add_heading(doc, '3.3  Optimized Inventory & Reduced Parts Waste', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'Maintenance departments without inventory visibility tend toward one of two costly '
        'extremes: over-stocking (capital tied up in parts that sit unused for years) or '
        'under-stocking (last-minute emergency purchases at premium prices with expedited '
        'shipping). AIT CMMS eliminates both problems.'
    )
    for item in [
        'Minimum stock level alerts prevent stockouts before they cause production delays.',
        'Unit-cost tracking and consumption history reveal which parts are fast-moving vs. slow-moving, enabling right-sizing of safety stock.',
        'Supplier lead-time data allows just-in-time ordering rather than panic buying.',
        'CM parts-request tracking prevents technicians from over-ordering "just in case."',
        'Total MRO valuation reports give management full visibility into inventory carrying costs.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()
    doc.add_paragraph(
        'Industry benchmarks indicate that CMMS-driven inventory management typically reduces '
        'annual MRO spend by 10–20% through elimination of duplicate ordering, expired '
        'consumables, and emergency premium purchases. For a maintenance department spending '
        '$150,000–$300,000 per year on parts and consumables, this represents $15,000–$60,000 '
        'in annual savings.'
    )
    doc.add_paragraph()

    # 3.4 Zero Software Licensing Cost
    add_heading(doc, '3.4  Zero Recurring Software Licensing Cost', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'AIT CMMS 2.3.1 was developed internally and runs on open-source components '
        '(Python, PostgreSQL). There are no per-seat licenses, no annual subscription fees, '
        'and no vendor lock-in. The system\'s database was recently migrated from a paid '
        'cloud hosting service (Neon) to local PostgreSQL on existing hardware — eliminating '
        'that cost entirely.'
    )

    lic_rows = [
        ('Commercial CMMS (e.g., eMaint, Fiix, UpKeep)', '$150–$500/user/year',  '$9,000–$30,000/year (est. 60 users)', 'Ongoing'),
        ('Previous cloud database (Neon)',                '$19–$69/month',         '$228–$828/year',                      'Ongoing'),
        ('AIT CMMS 2.3.1 (internal + local PostgreSQL)', '$0/year',               '$0/year',                             'None'),
    ]
    tbl_lic = doc.add_table(rows=1, cols=4)
    tbl_lic.style = 'Table Grid'
    lic_headers = ['Solution', 'Per-User Cost', 'Total Annual Cost', 'Licensing Risk']
    hdr_l = tbl_lic.rows[0].cells
    for i, h in enumerate(lic_headers):
        hdr_l[i].text = h
        hdr_l[i].paragraphs[0].runs[0].bold = True
        hdr_l[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(hdr_l[i], '1F3864')
        hdr_l[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, row_data in enumerate(lic_rows):
        row = tbl_lic.add_row().cells
        for j, val in enumerate(row_data):
            row[j].text = val
            row[j].paragraphs[0].runs[0].font.size = Pt(10)
            bg = 'E8F5E9' if i % 2 == 0 else 'FFFFFF'
            set_cell_bg(row[j], bg)
    doc.add_paragraph()
    doc.add_paragraph(
        'By avoiding commercial CMMS licensing fees, AIT saves an estimated $9,000–$30,000 '
        'per year — funds that remain available for direct maintenance activities.'
    )
    doc.add_paragraph()

    # 3.5 KPI-Driven Continuous Improvement
    add_heading(doc, '3.5  KPI-Driven Continuous Improvement & Cost Accountability', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'The 17+ KPI framework embedded in AIT CMMS transforms maintenance from a cost center '
        'into a measurable, managed business function. When leadership can see quantified '
        'metrics — PM adherence, MTBF trends, cost per event, backlog ratios — they can make '
        'data-driven budget decisions and hold the maintenance function accountable for '
        'continuous improvement.'
    )
    for item in [
        'MTBF trending identifies which assets should be scheduled for replacement before '
        'they generate runaway repair costs — preventing a $500 repair from becoming a '
        '$15,000 overhaul.',
        'Technical Availability data quantifies the dollar value of uptime protected by '
        'maintenance investment, making the ROI of the maintenance budget visible to Finance.',
        'Work Order Backlog KPI provides early warning of resource gaps, enabling proactive '
        'staffing decisions rather than reactive overtime.',
        'Injury Frequency Rate (FR1) tracking demonstrates safety compliance, which directly '
        'impacts insurance premiums, regulatory standing, and customer audit scores.',
        'Cost-per-event trending allows year-over-year benchmarking and drives targeted '
        'efficiency initiatives.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # 3.6 Compliance & Audit Risk Reduction
    add_heading(doc, '3.6  Compliance, Audit, & Risk Cost Reduction', level=2, color=(70, 130, 180))
    doc.add_paragraph(
        'The full audit trail built into AIT CMMS — every PM completion, every work order '
        'action, every parts transaction, logged with user name and timestamp — provides '
        'irrefutable proof of maintenance compliance during regulatory inspections, customer '
        'quality audits, and internal reviews. The financial value of this capability is '
        'significant:'
    )
    for item in [
        'Regulatory fines for undocumented maintenance can range from $5,000 to $500,000+ '
        'depending on the standard (ISO, AS9100, OSHA, etc.).',
        'Failed customer audits can result in lost contracts worth multiples of the fine '
        'exposure.',
        'Legal liability in equipment-failure incidents is substantially reduced when '
        'documented evidence of a thorough maintenance program is available.',
        'Insurance underwriters reward verifiable preventive maintenance programs with '
        'lower premiums on equipment and liability policies.',
    ]:
        add_bullet(doc, item)
    doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════
    # CONSOLIDATED ROI SUMMARY
    # ══════════════════════════════════════════
    add_heading(doc, '4. Consolidated Financial Impact Summary', level=1, color=(31, 56, 100))

    doc.add_paragraph(
        'The table below consolidates the estimated annual financial impact of AIT CMMS 2.3.1 '
        'across all value categories. Figures represent conservative estimates based on '
        'industry benchmarks; actual results will vary based on AIT\'s specific asset base, '
        'staffing levels, and baseline performance.'
    )

    summary_rows = [
        ('Reduced emergency repair costs (downtime avoidance)',    '$35,000',   '$120,000'),
        ('Labor productivity recovery (admin time eliminated)',     '$18,000',   '$28,600'),
        ('MRO inventory optimization (reduced waste & over-spend)','$15,000',   '$60,000'),
        ('Software licensing cost avoidance (vs. commercial CMMS)','$9,000',    '$30,000'),
        ('Compliance/audit risk avoidance (regulatory fines)',      '$5,000',    '$50,000+'),
        ('TOTAL ESTIMATED ANNUAL VALUE',                           '$82,000',   '$288,600+'),
    ]
    tbl_sum = doc.add_table(rows=1, cols=3)
    tbl_sum.style = 'Table Grid'
    sum_hdrs = tbl_sum.rows[0].cells
    for i, h in enumerate(['Value Category', 'Conservative Estimate', 'Optimistic Estimate']):
        sum_hdrs[i].text = h
        sum_hdrs[i].paragraphs[0].runs[0].bold = True
        sum_hdrs[i].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(sum_hdrs[i], '1F3864')
        sum_hdrs[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        sum_hdrs[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (cat, low, high) in enumerate(summary_rows):
        row = tbl_sum.add_row().cells
        row[0].text = cat
        row[1].text = low
        row[2].text = high
        is_total = (i == len(summary_rows) - 1)
        bg = '1F3864' if is_total else ('E2EFDA' if i % 2 == 0 else 'FFFFFF')
        txt_color = RGBColor(255, 255, 255) if is_total else RGBColor(0, 0, 0)
        for cell in row:
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].bold = is_total
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.color.rgb = txt_color
    doc.add_paragraph()

    doc.add_paragraph(
        'Note: These estimates reflect value generated, not necessarily cash savings that '
        'appear as line-item reductions in a budget. Many benefits (avoided downtime, avoided '
        'fines, faster labor) represent cost avoidance — expenditures that would have occurred '
        'without AIT CMMS but will not occur with it.'
    ).italic = True
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # STRATEGIC ADVANTAGES
    # ══════════════════════════════════════════
    add_heading(doc, '5. Strategic & Operational Advantages Beyond Cost', level=1, color=(31, 56, 100))

    strategic = [
        ('Data Ownership & No Vendor Risk',
         'All data resides on AIT-controlled servers. There is no risk of a vendor '
         'increasing prices, discontinuing a product, or suffering a data breach that '
         'exposes AIT\'s operational information.'),
        ('Scalability at No Additional Cost',
         'Adding new assets, users, or PM schedules costs nothing in licensing. As AIT '
         'grows, the system grows with it without triggering per-seat billing.'),
        ('Institutional Knowledge Capture',
         'Every repair note, parts-consumption record, and PM observation is stored '
         'permanently. When experienced technicians leave, their accumulated knowledge '
         'remains in the system for the next generation.'),
        ('Continuous In-House Development',
         'Because the system is owned and developed internally, new features can be '
         'added rapidly to respond to changing operational needs — no waiting for a '
         'vendor\'s release cycle or paying for custom development.'),
        ('Competitive Differentiation',
         'A documented, data-driven maintenance program with published KPIs and full '
         'audit trails positions AIT as a mature, reliable partner in quality-sensitive '
         'customer relationships and competitive bids.'),
    ]
    for title, body in strategic:
        p = doc.add_paragraph()
        run_title = p.add_run(title + ': ')
        run_title.bold = True
        run_title.font.size = Pt(11)
        run_body = p.add_run(body)
        run_body.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    # ══════════════════════════════════════════
    # CONCLUSION
    # ══════════════════════════════════════════
    add_heading(doc, '6. Conclusion', level=1, color=(31, 56, 100))

    doc.add_paragraph(
        'AIT CMMS 2.3.1 is not simply a digital filing cabinet for maintenance records. It '
        'is an active operational platform that drives measurable improvements in equipment '
        'reliability, workforce productivity, inventory efficiency, and regulatory compliance '
        '— translating directly into conservatively $82,000 to over $288,000 in annual '
        'financial value for the AIT organization.'
    )
    doc.add_paragraph(
        'Unlike commercial alternatives that impose ongoing subscription fees, per-user '
        'licensing, and vendor dependency, AIT CMMS is fully owned by AIT, runs on '
        'open-source infrastructure at zero recurring software cost, and evolves continuously '
        'to meet AIT\'s specific operational requirements.'
    )
    doc.add_paragraph(
        'The return on investment is clear, the risks are low, and the strategic benefits '
        'compound over time as the system\'s historical dataset grows. AIT CMMS 2.3.1 '
        'represents one of the highest-value, lowest-cost technology investments available '
        'to AIT\'s maintenance and operations leadership.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer line
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        '─────────────────────────────────────────────────────────────────\n'
        'AIT CMMS 2.3.1  |  Confidential — For Internal Distribution Only  |  '
        + datetime.date.today().strftime('%B %Y')
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(120, 120, 120)
    footer_run.italic = True

    # Save
    output_path = '/home/user/AIT_CMMS2.3.1/AIT_CMMS_Program_Overview_and_Financial_Value.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    build_document()
